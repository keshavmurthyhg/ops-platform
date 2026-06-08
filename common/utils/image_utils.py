import os
import uuid
import tempfile

from PIL import Image as PILImage

# PDF
from reportlab.platypus import Image, Spacer

# Word
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# -----------------------------------
# PDF IMAGE HANDLER
# -----------------------------------
def add_images_pdf(elements, images):
    """
    Add images to PDF sections.
    Supports:
    - file paths
    - raw bytes
    """

    if not images:
        return

    for img_data in images:
        try:
            temp_path = os.path.join(
                tempfile.gettempdir(),
                f"{uuid.uuid4()}.png"
            )

            # uploaded bytes
            if isinstance(img_data, bytes):
                with open(temp_path, "wb") as f:
                    f.write(img_data)

            # file path
            elif isinstance(img_data, str):
                if not os.path.exists(img_data):
                    continue

                temp_path = img_data

            else:
                continue

            pil_img = PILImage.open(temp_path)
            width, height = pil_img.size

            max_width = 450
            max_height = 500

            ratio = min(
                max_width / width,
                max_height / height
            )

            new_width = width * ratio
            new_height = height * ratio

            elements.append(
                Image(
                    temp_path,
                    width=new_width,
                    height=new_height
                )
            )

            # allow image splitting safely
            img.hAlign = "CENTER"

            elements.append(img)

            # smaller spacer
            elements.append(Spacer(1, 8))

        except Exception:
            continue


# -----------------------------------
# WORD IMAGE HANDLER
# -----------------------------------
def add_images_word(doc, images):
    """
    Add centered images to Word sections.
    """

    if not images:
        return

    for img_path in images:
        try:
            if not img_path:
                continue

            if not os.path.exists(img_path):
                continue

            # centered paragraph
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = p.add_run()

            run.add_picture(
                img_path,
                width=Inches(5.5)
            )

            # spacing after image
            doc.add_paragraph("")

        except Exception:
            continue