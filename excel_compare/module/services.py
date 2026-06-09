import os
import zipfile
import tempfile
import pandas as pd  # Added missing import
from datetime import datetime
from docx import Document
from openpyxl import load_workbook 
from .logic import load_all_sheets, diff_sheet, FILL_ADDED, FILL_MODIFIED, FILL_REMOVED, FONT_REMOVED

class ExcelExportService:
    @staticmethod
    def generate_bundle(path1, path2):
        # Explicitly use engine here as well for the high-level load
        wb1 = load_workbook(path1)
        wb2 = load_workbook(path2)
        
        s1 = load_all_sheets(path1)
        s2 = load_all_sheets(path2)
        all_sheets = list(dict.fromkeys(list(s1.keys()) + list(s2.keys())))

        # Document Report Compilation
        doc = Document()
        doc.add_heading("Excel Comparison Executive Summary", 0)
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%d-%b-%Y %H:%M:%S')}")

        for s_name in all_sheets:
            df1 = s1.get(s_name, pd.DataFrame()).reset_index(drop=True)
            df2 = s2.get(s_name, pd.DataFrame()).reset_index(drop=True)
            if df1.empty or df2.empty: continue
            dm1, dm2, summary = diff_sheet(df1, df2)

            if any(summary.values()):
                doc.add_heading(f"Sheet Name: {s_name}", level=1)
                doc.add_paragraph(f"• Modified Rows: {summary['modified']}")
                doc.add_paragraph(f"• Added Rows: {summary['added']}")
                doc.add_paragraph(f"• Removed Rows: {summary['removed']}")

            # Highlight New/Modified entries
            if s_name in wb2.sheetnames:
                ws2 = wb2[s_name]
                for r in range(len(df2)):
                    for c in range(len(df2.columns)):
                        if dm2.at[r, c] == "added": ws2.cell(r+1, c+1).fill = FILL_ADDED
                        elif dm2.at[r, c] == "modified": ws2.cell(r+1, c+1).fill = FILL_MODIFIED

            # Highlight Removed entries (Fixes f_mod typo)
            if s_name in wb1.sheetnames:
                ws1 = wb1[s_name]
                for r in range(len(df1)):
                    for c in range(len(df1.columns)):
                        if dm1.at[r, c] == "removed":
                            cell = ws1.cell(r+1, c+1)
                            cell.fill = FILL_REMOVED
                            cell.font = FONT_REMOVED

        t_dir = tempfile.mkdtemp()
        out1 = os.path.join(t_dir, f"OLD_{os.path.basename(path1)}")
        out2 = os.path.join(t_dir, f"NEW_{os.path.basename(path2)}")
        word_out = os.path.join(t_dir, "Comparison_Executive_Summary.docx")
        
        wb1.save(out1); wb2.save(out2); doc.save(word_out)

        z_name = f"ExcelCompare_{datetime.now().strftime('%d%b%Y')}.zip"
        z_path = os.path.join(t_dir, z_name)
        with zipfile.ZipFile(z_path, "w", zipfile.ZIP_DEFLATED) as z:
            z.write(out1, f"OLD_{os.path.basename(path1)}")
            z.write(out2, f"NEW_{os.path.basename(path2)}")
            z.write(word_out, "Comparison_Executive_Summary.docx")
            
        return z_path, z_name