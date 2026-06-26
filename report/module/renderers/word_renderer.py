from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from common.utils.image_utils import add_images_word
from report.module.layout.footer import apply_word_footer
from common.utils.links import apply_word_link, parse_ptc_cases
from common.utils.formatters import (
    format_date,
    set_cell_bg,
    safe_text
)
from common.utils.text_cleaner import (
    clean_text,
    format_description
)

# PTC support case URL — must match links.py get_url("ptc case", ...)
_PTC_CASE_URL = "https://support.ptc.com/appserver/cs/view/case.jsp?n={}"


# -----------------------------------
# CELL PADDING
# -----------------------------------
def set_cell_padding(
    cell,
    top=120,
    start=120,
    bottom=120,
    end=120
):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcMar = OxmlElement("w:tcMar")

    for margin_name, margin_value in [
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ]:
        node = OxmlElement(f"w:{margin_name}")
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)

    tcPr.append(tcMar)

def clean_azure_bug(value):
    """
    Prevent random invalid Azure bug numbers from showing in Word.
    Matches PDF behavior.
    """
    if value is None:
        return "-"

    value = str(value).strip()

    if value.lower() in ["nan", "none", "nat", ""]:
        return "-"

    return value


def apply_table_padding(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_padding(cell)


def generate_word_doc(
    data,
    root,
    l2,
    res,
    images,
    ppt_data=None
):
    doc = Document()

    # -----------------------------------
    # PAGE MARGINS
    # -----------------------------------
    section = doc.sections[0]
    
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    
    # Reduce left/right margins
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # -----------------------------------
    # TITLE
    # -----------------------------------
    title = doc.add_heading(
        "INCIDENT REPORT",
        0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # -----------------------------------
    # HEADER TABLE
    # Match PDF ratio: 100:160:100:160
    # Total width ≈ 6.5 inches
    # -----------------------------------
    table = doc.add_table(
        rows=4,
        cols=4
    )

    table.style = "Table Grid"
    table.autofit = False
    table.allow_autofit = False
    
    # Better proportional sizing
    column_widths = [
        Inches(1.0),   # label
        Inches(2.0),   # value
        Inches(1.5),   # label
        Inches(2.0)    # value
    ]

    for row in table.rows:
        for i, width in enumerate(column_widths):
            row.cells[i].width = width

    def fill(r, c, key, val):
        header_cell = table.rows[r].cells[c]
        value_cell = table.rows[r].cells[c + 1]

        # Header styling
        p = header_cell.paragraphs[0]
        run = p.add_run(key.upper())
        run.bold = True
        set_cell_bg(header_cell)

        cleaned_val = safe_text(val)

        if str(cleaned_val).lower() in [
            "nan",
            "nat",
            "none",
            ""
        ]:
            cleaned_val = "-"

        value_para = value_cell.paragraphs[0]

        apply_word_link(
            value_para,
            key,
            cleaned_val
        )

    fill(0, 0, "Incident", data.get("number"))
    fill(0, 2, "Created By", data.get("created_by"))

    fill(
        1,
        0,
        "Azure Bug",
        clean_azure_bug(data.get("azure_bug"))
    )
    
    fill(
        1,
        2,
        "Created Date",
        format_date(data.get("created_date"))
    )

    fill(2, 0, "PTC Case", data.get("ptc_case"))
    # ── Override PTC Case cell: multiple comma-separated cases, each individually
    #    hyperlinked. Alpha prefixes (C1234567) stripped for URL, kept for display.
    ptc_cell = table.rows[2].cells[1]
    ptc_para = ptc_cell.paragraphs[0]
    ptc_para.clear()
    ptc_cases = parse_ptc_cases(data.get("ptc_case"))  # returns [(display, num_id), ...]
    if not ptc_cases:
        ptc_para.add_run("-")
    else:
        for i, (display, num_id) in enumerate(ptc_cases):
            if i > 0:
                sep = ptc_para.add_run(", ")
                sep.font.size = Pt(10)
            try:
                url  = _PTC_CASE_URL.format(num_id)
                r_id = ptc_cell.part.relate_to(
                    url,
                    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                    is_external=True,
                )
                hl   = OxmlElement("w:hyperlink"); hl.set(qn("r:id"), r_id)
                wr   = OxmlElement("w:r")
                rPr  = OxmlElement("w:rPr")
                col  = OxmlElement("w:color"); col.set(qn("w:val"), "000000")
                u_el = OxmlElement("w:u");    u_el.set(qn("w:val"), "none")
                sz   = OxmlElement("w:sz");   sz.set(qn("w:val"),  "20")
                szCs = OxmlElement("w:szCs"); szCs.set(qn("w:val"), "20")
                rPr.extend([col, u_el, sz, szCs])
                wr.append(rPr)
                t_el = OxmlElement("w:t"); t_el.text = display
                wr.append(t_el); hl.append(wr); ptc_para._p.append(hl)
            except Exception:
                run = ptc_para.add_run(display); run.font.size = Pt(10)
    fill(2, 2, "Assigned To", data.get("assigned_to"))

    fill(3, 0, "Priority", data.get("priority"))
    fill(
        3,
        2,
        "Resolved Date",
        format_date(data.get("resolved_date"))
    )

    apply_table_padding(table)

    doc.add_paragraph("")

    # -----------------------------------
    # DESCRIPTION TABLE
    # Same total width as header table
    # -----------------------------------
    t2 = doc.add_table(
        rows=2,
        cols=2
    )
    
    t2.style = "Table Grid"
    t2.autofit = False
    t2.allow_autofit = False
    
    # Proportional widths
    desc_widths = [
        Inches(3.0),   # Short Description
        Inches(3.5)    # Description
    ]
    
    # Apply widths to all rows
    for row in t2.rows:
        for i, width in enumerate(desc_widths):
            row.cells[i].width = width
    
    # Header row
    headers = [
        "SHORT DESCRIPTION",
        "DESCRIPTION"
    ]
    
    for i, text in enumerate(headers):
        cell = t2.rows[0].cells[i]
        p = cell.paragraphs[0]
    
        # Clear default empty paragraph text
        p.clear()
    
        run = p.add_run(text)
        run.bold = True
    
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
        set_cell_bg(cell)
    
    # Data row
    t2.rows[1].cells[0].text = clean_text(
        safe_text(
            data.get("short_description")
        )
    )
    
    t2.rows[1].cells[1].text = clean_text(
        format_description(
            safe_text(
                data.get("description")
            )
        )
    )
    
    # Apply padding
    apply_table_padding(t2)
    
    doc.add_paragraph("")

    # -----------------------------------
    # RCA SECTIONS + IMAGES
    # -----------------------------------
    sections = [
        ("PROBLEM STATEMENT", root, images.get("problem", [])),
        ("ROOT CAUSE", l2, images.get("root", [])),
        ("RESOLUTION & RECOMMENDATION", res, images.get("resolution", []))
    ]

    for title, content, section_images in sections:

        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        content = safe_text(content)

        if str(content).lower() in ["nan", "nat", "none", ""]:
            content = "-"

        # Add text content
        for line in content.split("\n"):
            cleaned_line = clean_text(
                line.strip("- ").strip()
            )

            if cleaned_line:
                p = doc.add_paragraph(style="List Bullet")

                p.add_run(cleaned_line)

                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.first_line_indent = Inches(-0.18)
                p.paragraph_format.space_after = Pt(4)

        # Add uploaded images
        if section_images:
            add_images_word(
                doc,
                section_images
            )
        
    # -----------------------------------
    # REFERENCES TABLE — always starts on a new page
    # Headers: ALL CAPS  |  Col 0+1: center H+V  |  Col 2: left H, center V
    # Widths: Ref=1.8", Env=auto-fit(~0.9"), Link&Context=fills rest to 6.5"
    # -----------------------------------
    refs = data.get("references") or []
    if refs:
        doc.add_page_break()

        ref_heading = doc.add_heading("REFERENCES", level=1)
        ref_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph("")

        # Col widths: Ref=1.8", Env=0.9" (fits "ENVIRONMENT" + data), Link=3.8" → 6.5" total
        COL_WIDTHS = [Inches(1.5), Inches(1.2), Inches(3.8)]

        ref_table = doc.add_table(rows=1, cols=3)
        ref_table.style = "Table Grid"
        ref_table.autofit = False
        ref_table.allow_autofit = False

        def _set_cell_valign(cell, align="center"):
            """Set vertical alignment of a table cell."""
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            vAlign = OxmlElement("w:vAlign")
            vAlign.set(qn("w:val"), align)
            tcPr.append(vAlign)

        # ── Header row ──────────────────────────────────────────────────────
        hdr_labels = ["REFERENCE", "ENVIRONMENT", "LINK & CONTEXT"]
        hdr_cells  = ref_table.rows[0].cells
        for i, hdr_text in enumerate(hdr_labels):
            cell = hdr_cells[i]
            cell.width = COL_WIDTHS[i]
            set_cell_bg(cell)
            _set_cell_valign(cell, "center")
            p = cell.paragraphs[0]
            p.clear()
            # Col 0 + 1: center; Col 2: center (header only)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(hdr_text)
            run.bold = True
            run.font.size = Pt(10)

        # ── Data rows ────────────────────────────────────────────────────────
        for ref in refs:
            row_cells = ref_table.add_row().cells
            for i, cell in enumerate(row_cells):
                cell.width = COL_WIDTHS[i]

            # Col 0: Reference label — bold, center H+V
            p0 = row_cells[0].paragraphs[0]; p0.clear()
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_valign(row_cells[0], "center")
            run0 = p0.add_run(ref.get("label", ""))
            run0.bold = True
            run0.font.size = Pt(10)

            # Col 1: Environment — plain, center H+V
            p1 = row_cells[1].paragraphs[0]; p1.clear()
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_valign(row_cells[1], "center")
            run1 = p1.add_run(ref.get("environment") or "-")
            run1.font.size = Pt(10)

            # Col 2: URL + context — left H, center V
            _set_cell_valign(row_cells[2], "center")
            p2  = row_cells[2].paragraphs[0]; p2.clear()
            p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
            url = ref.get("url", "")
            ctx = ref.get("context", "")
            if url:
                try:
                    r_id = row_cells[2].part.relate_to(
                        url,
                        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                        is_external=True,
                    )
                    hl  = OxmlElement("w:hyperlink"); hl.set(qn("r:id"), r_id)
                    wr  = OxmlElement("w:r")
                    rPr = OxmlElement("w:rPr")
                    col_el = OxmlElement("w:color")
                    col_el.set(qn("w:val"), "000000")  # black
                    u_el2 = OxmlElement("w:u"); u_el2.set(qn("w:val"), "none")  # no underline
                    rPr.append(col_el); rPr.append(u_el2)
                    sz_el = OxmlElement("w:sz");  sz_el.set(qn("w:val"),  "20")  # 10pt
                    sz_cs = OxmlElement("w:szCs"); sz_cs.set(qn("w:val"), "20")
                    rPr.append(sz_el); rPr.append(sz_cs)
                    wr.append(rPr)
                    t_el = OxmlElement("w:t"); t_el.text = url
                    wr.append(t_el); hl.append(wr); p2._p.append(hl)
                except Exception:
                    run = p2.add_run(url); run.font.size = Pt(10)
            if ctx:
                p_ctx = row_cells[2].add_paragraph()
                p_ctx.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run_ctx = p_ctx.add_run(ctx)
                run_ctx.font.size      = Pt(9)
                run_ctx.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        apply_table_padding(ref_table)

    # -----------------------------------
    # FOOTER
    # -----------------------------------
    apply_word_footer(
        doc,
        data
    )

    # -----------------------------------
    # PPT SLIDES
    # -----------------------------------
    if ppt_data:
        try:
            from converter.module.ppt_slide_renderer import (
                render_ppt_slides_to_images
            )

            doc.add_page_break()
            doc.add_heading(
                "PPT Slides",
                level=1
            )

            slide_images = render_ppt_slides_to_images(
                ppt_data
            )

            if not slide_images:
                doc.add_paragraph(
                    "No slide images found in PPT."
                )
            else:
                for img in slide_images[1:]:
                    doc.add_page_break()
                    doc.add_picture(
                        img,
                        width=Inches(6.5)
                    )

        except Exception as e:
            doc.add_paragraph(
                f"Unable to attach PPT slides: {str(e)}"
            )

    # -----------------------------------
    # SAVE
    # -----------------------------------
    from io import BytesIO

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()
