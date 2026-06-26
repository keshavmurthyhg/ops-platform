from docx import Document

def extract_doc_content(file):
    doc = Document(file)

    paragraphs = []
    tables = []

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Extract tables
    for table_index, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)

        tables.append({
            "table_index": table_index + 1,
            "data": table_data
        })

    return {
        "paragraphs": paragraphs,
        "tables": tables
    }
