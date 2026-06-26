import difflib
import html
import streamlit.components.v1 as components
from docx import Document


# --------------------------------------------------
# Extract document content
# --------------------------------------------------
def extract_doc_content(doc_file):
    doc = Document(doc_file)
    content = []

    # -------------------------
    # Paragraphs
    # -------------------------
    for para in doc.paragraphs:
        text = para.text.strip()

        if text:
            content.append(text)

    # -------------------------
    # Tables
    # -------------------------
    for table_index, table in enumerate(doc.tables):
        content.append(
            f"[TABLE-{table_index+1}]"
        )

        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip()
                for cell in row.cells
            )

            content.append(
                f"[TABLE] {row_text}"
            )

    # -------------------------
    # Images
    # -------------------------
    image_count = 0

    for rel in doc.part.rels.values():
        try:
            if rel.is_external:
                continue

            if "image" in rel.target_ref.lower():
                image_count += 1

        except Exception:
            continue

    if image_count:
        content.append(
            f"[IMAGES FOUND: {image_count}]"
        )

    return content


# --------------------------------------------------
# Build row html
# --------------------------------------------------
def build_row(text, css_class):
    safe_text = html.escape(text)

    return f"""
    <div
        class="line {css_class}"
        title="{safe_text}"
    >
        {safe_text}
    </div>
    """


# --------------------------------------------------
# Generate aligned diff rows
# --------------------------------------------------
def generate_aligned_diff_rows(
    old_lines,
    new_lines
):
    matcher = difflib.SequenceMatcher(
        None,
        old_lines,
        new_lines
    )

    old_rows = []
    new_rows = []

    for opcode, i1, i2, j1, j2 in matcher.get_opcodes():

        # --------------------------------
        # Equal rows
        # --------------------------------
        if opcode == "equal":
            max_len = max(
                i2 - i1,
                j2 - j1
            )

            for idx in range(max_len):

                old_line = (
                    old_lines[i1 + idx]
                    if (i1 + idx) < i2
                    else ""
                )

                new_line = (
                    new_lines[j1 + idx]
                    if (j1 + idx) < j2
                    else ""
                )

                old_rows.append(
                    build_row(
                        old_line,
                        "normal"
                        if old_line
                        else "blank"
                    )
                )

                new_rows.append(
                    build_row(
                        new_line,
                        "normal"
                        if new_line
                        else "blank"
                    )
                )

        # --------------------------------
        # Deleted rows
        # --------------------------------
        elif opcode == "delete":
            deleted_lines = old_lines[i1:i2]

            for line in deleted_lines:
                old_rows.append(
                    build_row(
                        f"❌ Removed: {line}",
                        "removed"
                    )
                )

                new_rows.append(
                    build_row(
                        "",
                        "blank"
                    )
                )

        # --------------------------------
        # Added rows
        # --------------------------------
        elif opcode == "insert":
            inserted_lines = new_lines[j1:j2]

            for line in inserted_lines:
                old_rows.append(
                    build_row(
                        "",
                        "blank"
                    )
                )

                new_rows.append(
                    build_row(
                        f"➕ Added: {line}",
                        "added"
                    )
                )

        # --------------------------------
        # Replaced rows
        # --------------------------------
        elif opcode == "replace":
            old_chunk = old_lines[i1:i2]
            new_chunk = new_lines[j1:j2]

            max_len = max(
                len(old_chunk),
                len(new_chunk)
            )

            for idx in range(max_len):

                old_line = (
                    old_chunk[idx]
                    if idx < len(old_chunk)
                    else ""
                )

                new_line = (
                    new_chunk[idx]
                    if idx < len(new_chunk)
                    else ""
                )

                old_rows.append(
                    build_row(
                        f"🔄 Replaced: {old_line}"
                        if old_line else "",
                        "updated"
                        if old_line
                        else "blank"
                    )
                )

                new_rows.append(
                    build_row(
                        f"🔄 Updated: {new_line}"
                        if new_line else "",
                        "updated"
                        if new_line
                        else "blank"
                    )
                )

    return (
        "".join(old_rows),
        "".join(new_rows)
    )


# --------------------------------------------------
# Render synced preview
# --------------------------------------------------
def render_synced_preview(
    old_html,
    new_html
):
    combined_html = f"""
    <html>
    <head>
    <style>

        body {{
            margin:0;
            padding:0;
            font-family:Arial;
            overflow:hidden;
        }}

        .container {{
            display:flex;
            width:100%;
            height:360px;   /* show ~15 rows */
            border:1px solid #ccc;
            overflow:hidden;
        }}

        .pane {{
            width:50%;
            height:100%;
            overflow-y:auto;
            overflow-x:hidden;
            border-right:1px solid #ddd;
            font-family:Consolas, monospace;
        }}

        .line {{
            height:24px;
            line-height:24px;
            padding:0 6px;
            margin:0;
            font-size:12px;
            white-space:nowrap;
            overflow:hidden;
            text-overflow:ellipsis;
            border-radius:3px;
            box-sizing:border-box;
        }}

        .normal {{
            background:white;
        }}

        .removed {{
            background:#ffd6d6;
        }}

        .added {{
            background:#d6f5d6;
        }}

        .updated {{
            background:#fff2cc;
        }}

        .blank {{
            background:white;
        }}

        /* cleaner scrollbar */
        .pane::-webkit-scrollbar {{
            width:8px;
        }}

        .pane::-webkit-scrollbar-thumb {{
            background:#b5b5b5;
            border-radius:10px;
        }}

        .pane::-webkit-scrollbar-track {{
            background:#f5f5f5;
        }}

    </style>
    </head>

    <body>

        <div class="container">

            <div
                class="pane"
                id="leftPane"
            >
                {old_html}
            </div>

            <div
                class="pane"
                id="rightPane"
            >
                {new_html}
            </div>

        </div>

        <script>
            const left =
                document.getElementById("leftPane");

            const right =
                document.getElementById("rightPane");

            let syncing = false;

            left.addEventListener(
                "scroll",
                function() {{
                    if (!syncing) {{
                        syncing = true;
                        right.scrollTop = left.scrollTop;
                        syncing = false;
                    }}
                }}
            );

            right.addEventListener(
                "scroll",
                function() {{
                    if (!syncing) {{
                        syncing = true;
                        left.scrollTop = right.scrollTop;
                        syncing = false;
                    }}
                }}
            );
        </script>

    </body>
    </html>
    """

    components.html(
        combined_html,
        height=380,   # slightly larger than container
        scrolling=False
    )
