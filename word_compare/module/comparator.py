from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from copy import deepcopy
import hashlib
import difflib


# ------------------------------------------
# Highlight helpers
# ------------------------------------------
def highlight_run(run, color="yellow"):

    # Replaced/Updated
    if color == "yellow":
        run.font.highlight_color = (
            WD_COLOR_INDEX.YELLOW
        )

    # Added
    elif color == "green":
        run.font.highlight_color = (
            WD_COLOR_INDEX.BRIGHT_GREEN
        )

    # Deleted
    elif color == "red":
        run.font.color.rgb = RGBColor(
            255,
            255,
            255
        )

        shading_elm = parse_xml(
            r'<w:shd {} w:fill="FF0000"/>'.format(
                nsdecls("w")
            )
        )

        run._r.get_or_add_rPr().append(
            shading_elm
        )


# ------------------------------------------
# Compare paragraph text
# ------------------------------------------
def compare_paragraphs(
    old_doc,
    new_doc,
    old_output_doc,
    new_output_doc
):
    old_paras = old_doc.paragraphs
    new_paras = new_doc.paragraphs

    max_len = max(
        len(old_paras),
        len(new_paras)
    )

    for i in range(max_len):

        # -------------------------
        # Entire paragraph added
        # -------------------------
        if i >= len(old_paras):
            if i < len(new_output_doc.paragraphs):
                for run in new_output_doc.paragraphs[i].runs:
                    highlight_run(run, "green")
            continue

        # -------------------------
        # Entire paragraph deleted
        # -------------------------
        if i >= len(new_paras):
            if i < len(old_output_doc.paragraphs):
                for run in old_output_doc.paragraphs[i].runs:
                    highlight_run(run, "red")
            continue

        old_text = old_paras[i].text.strip()
        new_text = new_paras[i].text.strip()

        if old_text == new_text:
            continue

        old_words = old_text.split()
        new_words = new_text.split()

        matcher = difflib.SequenceMatcher(
            None,
            old_words,
            new_words
        )

        operations = matcher.get_opcodes()

        # -------------------------
        # OLD FILE
        # Only true DELETE ops
        # -------------------------
        has_actual_delete = any(
            tag == "delete"
            for tag, *_ in operations
        )

        if has_actual_delete:
            old_output_doc.paragraphs[i].clear()

            for tag, a1, a2, b1, b2 in operations:
                segment = " ".join(
                    old_words[a1:a2]
                )

                if not segment:
                    continue

                run = old_output_doc.paragraphs[i].add_run(
                    segment + " "
                )

                if tag == "delete":
                    highlight_run(
                        run,
                        "red"
                    )

        # -------------------------
        # NEW FILE
        # Added → green
        # Updated → yellow
        # -------------------------
        has_insert = any(
            tag == "insert"
            for tag, *_ in operations
        )

        has_replace = any(
            tag == "replace"
            for tag, *_ in operations
        )

        if has_insert or has_replace:
            for run in new_output_doc.paragraphs[i].runs:

                if has_insert and not has_replace:
                    highlight_run(
                        run,
                        "green"
                    )

                elif has_replace:
                    highlight_run(
                        run,
                        "yellow"
                    )
# ------------------------------------------
# Compare tables
# ------------------------------------------
def compare_tables(
    old_doc,
    new_doc,
    old_output_doc,
    new_output_doc
):
    old_tables = old_doc.tables
    new_tables = new_doc.tables

    for t_idx in range(
        min(
            len(old_tables),
            len(new_tables)
        )
    ):
        old_table = old_tables[t_idx]
        new_table = new_tables[t_idx]

        old_output_table = old_output_doc.tables[t_idx]
        new_output_table = new_output_doc.tables[t_idx]

        for r_idx in range(
            min(
                len(old_table.rows),
                len(new_table.rows)
            )
        ):
            for c_idx in range(
                min(
                    len(
                        old_table.rows[r_idx].cells
                    ),
                    len(
                        new_table.rows[r_idx].cells
                    )
                )
            ):
                old_text = (
                    old_table.rows[r_idx]
                    .cells[c_idx]
                    .text
                )

                new_text = (
                    new_table.rows[r_idx]
                    .cells[c_idx]
                    .text
                )

                if old_text == new_text:
                    continue

                # Only deleted table content
                if old_text and not new_text:
                    old_cell = (
                        old_output_table
                        .rows[r_idx]
                        .cells[c_idx]
                    )

                    for para in old_cell.paragraphs:
                        for run in para.runs:
                            highlight_run(
                                run,
                                "red"
                            )

                # New table → updated
                new_cell = (
                    new_output_table
                    .rows[r_idx]
                    .cells[c_idx]
                )

                for para in new_cell.paragraphs:
                    for run in para.runs:
                        highlight_run(
                            run,
                            "yellow"
                        )


# ------------------------------------------
# Image hashes
# ------------------------------------------
def get_image_hashes(doc):
    hashes = []

    for rel in doc.part.rels.values():
        try:
            if rel.is_external:
                continue

            if "image" in rel.target_ref.lower():
                img_data = rel.target_part.blob
                img_hash = hashlib.md5(
                    img_data
                ).hexdigest()

                hashes.append(
                    img_hash
                )

        except Exception:
            continue

    return hashes


# ------------------------------------------
# Compare images
# ------------------------------------------
def compare_images(
    old_doc,
    new_doc,
    old_output_doc,
    new_output_doc
):
    old_images = set(
        get_image_hashes(old_doc)
    )

    new_images = set(
        get_image_hashes(new_doc)
    )

    added = new_images - old_images
    removed = old_images - new_images

    if added:
        p = new_output_doc.add_paragraph()

        run = p.add_run(
            f"Images Added: {len(added)}"
        )

        highlight_run(
            run,
            "green"
        )

    if removed:
        p = old_output_doc.add_paragraph()

        run = p.add_run(
            f"Images Removed: {len(removed)}"
        )

        highlight_run(
            run,
            "red"
        )


# ------------------------------------------
# Main compare
# ------------------------------------------
def compare_documents(
    old_file,
    new_file,
    old_output_path,
    new_output_path
):
    old_doc = Document(old_file)
    new_doc = Document(new_file)

    old_output_doc = deepcopy(
        old_doc
    )

    new_output_doc = deepcopy(
        new_doc
    )

    compare_paragraphs(
        old_doc,
        new_doc,
        old_output_doc,
        new_output_doc
    )

    compare_tables(
        old_doc,
        new_doc,
        old_output_doc,
        new_output_doc
    )

    compare_images(
        old_doc,
        new_doc,
        old_output_doc,
        new_output_doc
    )

    old_output_doc.save(
        old_output_path
    )

    new_output_doc.save(
        new_output_path
    )
