"""
word_compare/module/logic.py
Sections-based comparison engine for Word documents.
A "section" is detected by heading paragraphs (style starts with 'Heading')
or by a configurable chunk size for non-headed documents.
Returns section-keyed diff data analogous to Excel sheet data.
"""

import os
import re
import base64
import difflib
import hashlib
import html
import zipfile
import tempfile
from copy import deepcopy
from datetime import datetime, timezone, timedelta

from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from openpyxl import Workbook
from openpyxl.styles import (
    PatternFill, Font, Alignment, Border, Side, GradientFill
)
from openpyxl.utils import get_column_letter


# ─────────────────────────────────────────
# Highlight helpers
# ─────────────────────────────────────────

def _highlight_run(run, color="yellow"):
    if color == "yellow":
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    elif color == "green":
        run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
    elif color == "red":
        run.font.color.rgb = RGBColor(255, 255, 255)
        shading = parse_xml(r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls("w")))
        run._r.get_or_add_rPr().append(shading)


# ─────────────────────────────────────────
# Section detection & content extraction
# ─────────────────────────────────────────

def _is_heading(para):
    """True if the paragraph uses a Heading style."""
    style_name = (para.style.name or "").lower() if para.style else ""
    return style_name.startswith("heading")


def _extract_sections(doc_path):
    """
    Split document into named sections keyed by their heading text.
    Paragraphs before the first heading go into '__preamble__'.
    Tables are appended to the section that precedes them.
    Returns OrderedDict: {section_name: [line_text, ...]}
    """
    doc = Document(doc_path)
    sections = {}
    order = []

    current_key = "__preamble__"
    sections[current_key] = []
    order.append(current_key)

    # Build a set of paragraph indices that are inside tables
    # (doc.paragraphs includes table cell paragraphs in some builds, avoid double-count)
    para_list = doc.paragraphs  # top-level only in recent python-docx

    for para in para_list:
        text = para.text.strip()
        if not text:
            continue

        if _is_heading(para):
            # Use heading text as section key; deduplicate if repeated
            key = text
            if key in sections:
                key = f"{text} ({len([k for k in order if k.startswith(text)])})"
            sections[key] = [text]   # heading itself is the first line
            order.append(key)
            current_key = key
        else:
            sections[current_key].append(text)

    # Append table rows to the section they physically follow
    # We re-open doc to iterate body elements in order
    doc2 = Document(doc_path)
    from docx.oxml.ns import qn
    body = doc2.element.body
    current_key = "__preamble__"
    table_idx = 0

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            from docx.text.paragraph import Paragraph
            para = Paragraph(child, doc2)
            text = para.text.strip()
            if text and _is_heading(para):
                # find which key this heading maps to
                for k in order:
                    if k == text or k.startswith(text + " ("):
                        if sections.get(k) and sections[k][0] == text:
                            current_key = k
                            break
        elif tag == "tbl":
            table_idx += 1
            from docx.table import Table
            tbl = Table(child, doc2)
            sections[current_key].append(f"[TABLE {table_idx}]")
            for row in tbl.rows:
                row_text = " │ ".join(c.text.strip() for c in row.cells)
                if row_text.replace("│", "").strip():
                    sections[current_key].append(row_text)

    # Remove empty preamble
    if not sections.get("__preamble__"):
        order.remove("__preamble__")
        del sections["__preamble__"]

    return order, sections


def _extract_images_b64(doc_path):
    """Return list of base64-encoded images from docx."""
    doc = Document(doc_path)
    images = []
    seen = set()
    for rel in doc.part.rels.values():
        try:
            if rel.is_external:
                continue
            if "image" in rel.target_ref.lower():
                blob = rel.target_part.blob
                img_hash = hashlib.md5(blob).hexdigest()
                if img_hash in seen:
                    continue
                seen.add(img_hash)
                ext = rel.target_ref.rsplit(".", 1)[-1].lower()
                mime_map = {"jpg": "jpeg", "jpeg": "jpeg", "png": "png",
                            "gif": "gif", "bmp": "png", "wmf": "png"}
                mime = mime_map.get(ext, "png")
                b64 = base64.b64encode(blob).decode("utf-8")
                images.append({"mime": f"image/{mime}", "data": b64})
        except Exception:
            continue
    return images


# ─────────────────────────────────────────
# Diff row builder
# ─────────────────────────────────────────

def _build_diff_rows(old_lines, new_lines):
    matcher = difflib.SequenceMatcher(None, old_lines, new_lines, autojunk=False)
    old_rows, new_rows = [], []

    for opcode, i1, i2, j1, j2 in matcher.get_opcodes():
        if opcode == "equal":
            for k in range(i2 - i1):
                old_rows.append({"text": old_lines[i1 + k], "css": "normal"})
                new_rows.append({"text": new_lines[j1 + k], "css": "normal"})

        elif opcode == "delete":
            for line in old_lines[i1:i2]:
                old_rows.append({"text": line, "css": "removed"})
                new_rows.append({"text": "", "css": "blank"})

        elif opcode == "insert":
            for line in new_lines[j1:j2]:
                old_rows.append({"text": "", "css": "blank"})
                new_rows.append({"text": line, "css": "added"})

        elif opcode == "replace":
            old_chunk = old_lines[i1:i2]
            new_chunk = new_lines[j1:j2]
            max_len = max(len(old_chunk), len(new_chunk))
            for k in range(max_len):
                o = old_chunk[k] if k < len(old_chunk) else ""
                n = new_chunk[k] if k < len(new_chunk) else ""
                old_rows.append({"text": o, "css": "updated" if o else "blank"})
                new_rows.append({"text": n, "css": "updated" if n else "blank"})

    return old_rows, new_rows


def _section_stats(old_rows, new_rows):
    added   = sum(1 for r in new_rows if r["css"] == "added")
    removed = sum(1 for r in old_rows if r["css"] == "removed")
    modified = sum(1 for r in old_rows if r["css"] == "updated")
    return {"added": added, "removed": removed, "modified": modified}


# ─────────────────────────────────────────
# Public: run_compare
# ─────────────────────────────────────────

def run_compare(path1, path2):
    order1, secs1 = _extract_sections(path1)
    order2, secs2 = _extract_sections(path2)

    # Union of section keys, preserving order
    all_keys = list(dict.fromkeys(order1 + order2))

    section_data = {}
    t_added = t_removed = t_modified = 0

    for key in all_keys:
        old_lines = secs1.get(key, [])
        new_lines = secs2.get(key, [])
        if not old_lines and not new_lines:
            continue

        old_rows, new_rows = _build_diff_rows(old_lines, new_lines)
        stats = _section_stats(old_rows, new_rows)

        t_added    += stats["added"]
        t_removed  += stats["removed"]
        t_modified += stats["modified"]

        # Build change log for this section
        log = []
        max_len = max(len(old_rows), len(new_rows))
        for i in range(max_len):
            o = old_rows[i] if i < len(old_rows) else {"text": "", "css": "blank"}
            n = new_rows[i] if i < len(new_rows) else {"text": "", "css": "blank"}
            change_type = (
                "removed" if o["css"] == "removed"
                else "added" if n["css"] == "added"
                else "updated" if o["css"] == "updated" or n["css"] == "updated"
                else None
            )
            if change_type:
                log.append({
                    "section": key,
                    "old": o["text"],
                    "new": n["text"],
                    "type": change_type
                })

        section_data[key] = {
            "old_rows": old_rows,
            "new_rows": new_rows,
            "change_log": log,
            "added": stats["added"],
            "removed": stats["removed"],
            "modified": stats["modified"],
        }

    images_old = _extract_images_b64(path1)
    images_new = _extract_images_b64(path2)

    return {
        "sections": list(section_data.keys()),
        "section_data": section_data,
        "totals": {
            "added": t_added,
            "removed": t_removed,
            "modified": t_modified,
            "total": t_added + t_removed + t_modified,
        },
        "images_old": images_old,
        "images_new": images_new,
    }


# ─────────────────────────────────────────
# Highlighted .docx generation
# ─────────────────────────────────────────

def _compare_and_highlight(path1, path2, out_old, out_new):
    old_doc = Document(path1)
    new_doc = Document(path2)
    old_out = deepcopy(old_doc)
    new_out = deepcopy(new_doc)

    old_paras = old_doc.paragraphs
    new_paras = new_doc.paragraphs
    max_len = max(len(old_paras), len(new_paras))

    for i in range(max_len):
        if i >= len(old_paras):
            if i < len(new_out.paragraphs):
                for run in new_out.paragraphs[i].runs:
                    _highlight_run(run, "green")
            continue
        if i >= len(new_paras):
            if i < len(old_out.paragraphs):
                for run in old_out.paragraphs[i].runs:
                    _highlight_run(run, "red")
            continue

        old_text = old_paras[i].text.strip()
        new_text = new_paras[i].text.strip()
        if old_text == new_text:
            continue

        old_words = old_text.split()
        new_words = new_text.split()
        m = difflib.SequenceMatcher(None, old_words, new_words)
        ops = m.get_opcodes()

        if any(t == "delete" for t, *_ in ops):
            old_out.paragraphs[i].clear()
            for tag, a1, a2, b1, b2 in ops:
                seg = " ".join(old_words[a1:a2])
                if not seg:
                    continue
                run = old_out.paragraphs[i].add_run(seg + " ")
                if tag == "delete":
                    _highlight_run(run, "red")

        if any(t in ("insert", "replace") for t, *_ in ops):
            for run in new_out.paragraphs[i].runs:
                if any(t == "replace" for t, *_ in ops):
                    _highlight_run(run, "yellow")
                else:
                    _highlight_run(run, "green")

    for t_idx in range(min(len(old_doc.tables), len(new_doc.tables))):
        old_tbl = old_doc.tables[t_idx]
        new_tbl = new_doc.tables[t_idx]
        old_out_tbl = old_out.tables[t_idx]
        new_out_tbl = new_out.tables[t_idx]

        for r_idx in range(min(len(old_tbl.rows), len(new_tbl.rows))):
            for c_idx in range(min(len(old_tbl.rows[r_idx].cells),
                                   len(new_tbl.rows[r_idx].cells))):
                ot = old_tbl.rows[r_idx].cells[c_idx].text
                nt = new_tbl.rows[r_idx].cells[c_idx].text
                if ot == nt:
                    continue
                if ot and not nt:
                    for p in old_out_tbl.rows[r_idx].cells[c_idx].paragraphs:
                        for run in p.runs:
                            _highlight_run(run, "red")
                for p in new_out_tbl.rows[r_idx].cells[c_idx].paragraphs:
                    for run in p.runs:
                        _highlight_run(run, "yellow")

    old_out.save(out_old)
    new_out.save(out_new)


# ─────────────────────────────────────────
# Excel change log builder
# ─────────────────────────────────────────

# Colour constants
_HDR_FILL  = PatternFill("solid", fgColor="1E3A5F")
_ADD_FILL  = PatternFill("solid", fgColor="D6F5D6")
_REM_FILL  = PatternFill("solid", fgColor="FFD6D6")
_MOD_FILL  = PatternFill("solid", fgColor="FFF2CC")
_ALT_FILL  = PatternFill("solid", fgColor="F8FAFC")
_HDR_FONT  = Font(name="Calibri", bold=True, color="FFFFFF", size=11)
_BODY_FONT = Font(name="Calibri", size=10)
_BOLD_FONT = Font(name="Calibri", bold=True, size=10)

_thin = Side(style="thin", color="CBD5E1")
_BORDER = Border(left=_thin, right=_thin, top=_thin, bottom=_thin)

_TYPE_LABELS = {"added": "Added", "removed": "Removed", "updated": "Modified"}
_TYPE_FILLS  = {"added": _ADD_FILL, "removed": _REM_FILL, "updated": _MOD_FILL}
_TYPE_COLORS = {"added": "15803D", "removed": "B91C1C", "updated": "B45309"}


def _build_changelog_xlsx(section_data, section_keys, out_path,
                          file1_name, file2_name, totals):
    wb = Workbook()

    # ── Summary sheet ────────────────────────────────────────────
    ws_sum = wb.active
    ws_sum.title = "Summary"
    ws_sum.sheet_view.showGridLines = False
    ws_sum.column_dimensions["A"].width = 28
    ws_sum.column_dimensions["B"].width = 18

    # Title block
    ws_sum.merge_cells("A1:B1")
    ws_sum["A1"] = "Word Document Comparison — Change Log"
    ws_sum["A1"].font = Font(name="Calibri", bold=True, size=14, color="1E3A5F")
    ws_sum["A1"].fill = PatternFill("solid", fgColor="EFF6FF")
    ws_sum["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws_sum.row_dimensions[1].height = 32

    meta = [
        ("Base Document",     file1_name),
        ("Revised Document",  file2_name),
        ("Generated",         datetime.now(timezone(timedelta(hours=5, minutes=30)))
                                        .strftime("%d %b %Y  %H:%M IST")),
        ("", ""),
        ("Total Sections Compared", len(section_keys)),
        ("Lines Added",    totals.get("added", 0)),
        ("Lines Removed",  totals.get("removed", 0)),
        ("Lines Modified", totals.get("modified", 0)),
        ("Total Changes",  totals.get("total", 0)),
    ]
    for r_idx, (label, value) in enumerate(meta, start=2):
        ws_sum.row_dimensions[r_idx].height = 20
        a = ws_sum.cell(r_idx, 1, label)
        b = ws_sum.cell(r_idx, 2, value)
        a.font = _BOLD_FONT
        b.font = _BODY_FONT
        a.alignment = Alignment(vertical="center")
        b.alignment = Alignment(vertical="center")
        if label in ("Lines Added",):
            b.font = Font(name="Calibri", size=10, color="15803D", bold=True)
        elif label == "Lines Removed":
            b.font = Font(name="Calibri", size=10, color="B91C1C", bold=True)
        elif label == "Lines Modified":
            b.font = Font(name="Calibri", size=10, color="B45309", bold=True)
        elif label == "Total Changes":
            b.font = Font(name="Calibri", size=10, bold=True)

    # ── Per-section breakdown on summary ─────────────────────────
    ws_sum.row_dimensions[12].height = 20
    hdr_row = 13
    for col, txt in enumerate(["Section", "Added", "Removed", "Modified", "Total"], 1):
        c = ws_sum.cell(hdr_row, col, txt)
        c.fill = _HDR_FILL
        c.font = _HDR_FONT
        c.alignment = Alignment(horizontal="center", vertical="center")
        c.border = _BORDER
    ws_sum.column_dimensions["C"].width = 12
    ws_sum.column_dimensions["D"].width = 12
    ws_sum.column_dimensions["E"].width = 12
    ws_sum.row_dimensions[hdr_row].height = 22

    for i, key in enumerate(section_keys):
        sd  = section_data.get(key, {})
        a   = sd.get("added", 0)
        rem = sd.get("removed", 0)
        mod = sd.get("modified", 0)
        tot = a + rem + mod
        if tot == 0:
            continue
        r = hdr_row + 1 + i
        ws_sum.row_dimensions[r].height = 18
        row_fill = _ALT_FILL if i % 2 else PatternFill("solid", fgColor="FFFFFF")
        sec_short = key if len(key) <= 60 else key[:58] + "…"
        vals = [sec_short, a, rem, mod, tot]
        for col, v in enumerate(vals, 1):
            c = ws_sum.cell(r, col, v)
            c.fill = row_fill
            c.font = _BODY_FONT
            c.border = _BORDER
            c.alignment = Alignment(vertical="center",
                                    wrap_text=(col == 1))

    # ── Change Log sheet ─────────────────────────────────────────
    ws_log = wb.create_sheet("Change Log")
    ws_log.sheet_view.showGridLines = False
    ws_log.freeze_panes = "A2"

    headers = ["#", "Section", "Change Type", "Base Content", "Revised Content"]
    col_widths = [6, 32, 14, 55, 55]
    for col, (h, w) in enumerate(zip(headers, col_widths), 1):
        ws_log.column_dimensions[get_column_letter(col)].width = w
        c = ws_log.cell(1, col, h)
        c.fill = _HDR_FILL
        c.font = _HDR_FONT
        c.alignment = Alignment(horizontal="center", vertical="center")
        c.border = _BORDER
    ws_log.row_dimensions[1].height = 24

    row_num = 2
    entry_idx = 1
    for key in section_keys:
        sd = section_data.get(key, {})
        log = sd.get("change_log", [])
        for entry in log:
            change_type = entry.get("type", "")
            fill = _TYPE_FILLS.get(change_type, PatternFill())
            label = _TYPE_LABELS.get(change_type, change_type.capitalize())

            ws_log.row_dimensions[row_num].height = 40
            vals = [
                entry_idx,
                key,
                label,
                entry.get("old", ""),
                entry.get("new", ""),
            ]
            for col, v in enumerate(vals, 1):
                c = ws_log.cell(row_num, col, v)
                c.fill = fill
                c.font = Font(
                    name="Calibri", size=10,
                    color=_TYPE_COLORS.get(change_type, "1E293B"),
                    bold=(col == 3)
                )
                c.alignment = Alignment(vertical="top", wrap_text=True)
                c.border = _BORDER

            row_num   += 1
            entry_idx += 1

    wb.save(out_path)


# ─────────────────────────────────────────
# Highlighted .docx + changelog bundle
# ─────────────────────────────────────────

def generate_highlighted_bundle(path1, path2, name1, name2,
                                section_data=None, section_keys=None, totals=None):
    ist      = timezone(timedelta(hours=5, minutes=30))
    date_str = datetime.now(ist).strftime("%d%b%Y")

    base1      = os.path.splitext(name1)[0]
    base2      = os.path.splitext(name2)[0]
    out_name1  = f"{base1}_Diff-Highlighted_{date_str}.docx"
    out_name2  = f"{base2}_Diff-Highlighted_{date_str}.docx"
    log_name   = f"Word-Compare_ChangeLog_{date_str}.xlsx"
    zip_name   = f"Word-Compare_{date_str}.zip"

    with tempfile.TemporaryDirectory() as tmp:
        out1     = os.path.join(tmp, out_name1)
        out2     = os.path.join(tmp, out_name2)
        log_path = os.path.join(tmp, log_name)

        _compare_and_highlight(path1, path2, out1, out2)

        # Build change log only if data is available
        if section_data and section_keys:
            _build_changelog_xlsx(
                section_data, section_keys, log_path,
                name1, name2, totals or {}
            )

        zip_path = os.path.join(tempfile.gettempdir(), zip_name)
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.write(out1, out_name1)
            zf.write(out2, out_name2)
            if os.path.exists(log_path):
                zf.write(log_path, log_name)

    return zip_path, zip_name
