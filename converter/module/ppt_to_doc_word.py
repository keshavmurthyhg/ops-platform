import os
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT

from converter.module.ppt_extractor import extract_ppt_content
from converter.module.ppt_slide_renderer import (
    render_ppt_slides_to_images
)

from common.logger import setup_logger

logger = setup_logger("ppt_to_doc")

# ----------------------------------------
# Add extracted slide images into Word
# ----------------------------------------
def add_images_to_doc(doc, image_paths):
    """
    Insert exported PPT slide images
    into Word document using landscape layout
    """

    for index, img_path in enumerate(image_paths, start=1):
        try:
            if not os.path.exists(img_path):
                logger.warning(f"Image not found: {img_path}")
                continue

            # create new page for every slide
            doc.add_page_break()

            section = doc.sections[-1]

            # landscape orientation
            section.orientation = WD_ORIENT.LANDSCAPE

            section.page_width = Inches(11.69)
            section.page_height = Inches(8.27)

            # reduce margins
            section.left_margin = Inches(0.3)
            section.right_margin = Inches(0.3)
            section.top_margin = Inches(0.3)
            section.bottom_margin = Inches(0.3)

            # optional slide heading
            doc.add_heading(
                f"Slide {index}",
                level=2
            )

            # insert full slide image
            doc.add_picture(
                img_path,
                width=Inches(10.5)
            )

            logger.info(f"Inserted image: {img_path}")

        except Exception as e:
            logger.error(
                f"Failed inserting {img_path}: {str(e)}"
            )


# ----------------------------------------
# Convert PPT → Word
# ----------------------------------------
def convert_ppt_to_doc(ppt_path, output_docx):
    """
    Converts PPT into Word document:
    - extracts text
    - exports rendered slides
    - inserts slides into Word
    """

    try:
        logger.info(f"PPT to DOC started: {ppt_path}")
        doc = Document()

        # --------------------------------
        # Extract textual content
        # --------------------------------
        extracted_text = extract_ppt_content(ppt_path)

        if extracted_text:
            logger.info(
                f"Extracted {len(extracted_text)} text blocks"
            )

            doc.add_heading(
                "PPT Content",
                level=1
            )

            for text in extracted_text:
                if text and text.strip():
                    doc.add_paragraph(text)

        else:
            logger.warning("No text extracted from PPT")

        # --------------------------------
        # Export full rendered slides
        # --------------------------------
        image_paths = render_ppt_slides_to_images(
            ppt_path
        )

        logger.info(
            f"Extracted {len(image_paths)} slide images"
        )

        # --------------------------------
        # Add slides to Word
        # --------------------------------
        if image_paths:
            doc.add_page_break()

            doc.add_heading(
                "PPT Slides",
                level=1
            )

            add_images_to_doc(
                doc,
                image_paths
            )

        else:
            logger.warning("No slide images found")

            doc.add_page_break()

            doc.add_heading(
                "PPT Slides",
                level=1
            )

            doc.add_paragraph(
                "No slide images found in PPT."
            )

        # --------------------------------
        # Save final document
        # --------------------------------
        doc.save(output_docx)

        logger.info(
            f"Word document created successfully: {output_docx}"
        )

        return output_docx

    except Exception as e:
        logger.error(
            f"PPT conversion failed: {str(e)}"
        )
        raise


# ----------------------------------------
# Wrapper
# ----------------------------------------
def ppt_to_word(
    ppt_path,
    output_docx
):
    logger.info("ppt_to_word wrapper triggered")

    return convert_ppt_to_doc(
        ppt_path,
        output_docx
    )